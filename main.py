from docx import Document
import random
import re

# class ReportFormatter:
#     def __init__(self, file_path, color="black", font="Arial", 
#                  page_background="#FFFFFF", report_background="#F5F5F5", 
#                  font_size=14, line_height=1.6, page_width=1000, page_height=700):
#         self.file_path = file_path
#         self.color = color
#         self.font = font
#         self.page_background = page_background
#         self.report_background = report_background
#         self.font_size = font_size
#         self.line_height = line_height
#         self.page_width = page_width
#         self.page_height = page_height
#         self.content = ""
#         self.html_content = ""
        
#     def read_docx(self):
#         doc = Document(self.file_path)
#         content = ""
#         for paragraph in doc.paragraphs:
#             paragraph_text = paragraph.text.strip()
#             if re.match(r'^(Seção|SEÇÃO)\s+\d+|^(Introdução|Objetivo|Conclusão|Referências)', paragraph_text):
#                 content += f"<section-break> <div class='section-title'>{paragraph_text}</div>\n"
#             else:
#                 content += paragraph_text + "\n"
#         self.content = content.strip()
#         return self.content

#     def split_content_into_pages(self):
#         chars_per_line = int(self.page_width / (self.font_size * 0.85))  
#         line_height_px = self.font_size * self.line_height
#         lines_per_page = int(self.page_height / line_height_px)
#         max_chars_per_page = chars_per_line * lines_per_page

#         words = self.content.split()
#         pages = []
#         current_page = ""
#         current_chars = 0

#         for word in words:
#             print(f'THREADING WORD BY REPORT PROCESS ({word}): {random.getrandbits(128)}')

#             if current_chars + len(word) + 1 > max_chars_per_page:
#                 pages.append(current_page.strip())
#                 current_page = ""
#                 current_chars = 0
#             elif re.match(r'<section-break>', word):
#                 pages.append(current_page.strip())
#                 current_page = ""
#                 current_chars = 0

#             current_page += word + " "
#             current_chars += len(word) + 1

#         if current_page:
#             pages.append(current_page.strip())

#         return pages
  
#     def format_content_to_html(self):
#         pages = self.split_content_into_pages()
#         formatted_content = ""

#         for page_number, page_content in enumerate(pages, start=1):
#             formatted_content += f"""
#             <div class="page">
#                 {page_content.replace("\n", "<br>")}
#                 <div class="page-number">Página {page_number}</div>
#             </div>
#             """

#         self.html_content = f"""
#         <!DOCTYPE html>
#         <html lang="en">
#         <head>
#             <meta charset="UTF-8">
#             <meta name="viewport" content="width=device-width, initial-scale=1.0">
#             <title>Formatted Report</title>
#             <style>
#                 body {{
#                     font-family: {self.font};
#                     background-color: {self.report_background};
#                     display: flex;
#                     justify-content: center;
#                     padding: 20px;
#                     flex-wrap: wrap;
#                 }}
#                 .page {{
#                     background-color: {self.page_background};
#                     color: {self.color};
#                     width: {self.page_width}px;
#                     height: {self.page_height}px;
#                     padding: 70px 30px;
#                     margin: 15px;
#                     box-shadow: 2px 2px 5px rgba(0, 0, 0, 0.1);
#                     overflow: hidden;
#                     position: relative;
#                     font-size: {self.font_size}px;
#                     line-height: {self.line_height};
#                     box-sizing: border-box;
#                     position: relative;
#                 }}
#                 .section-title {{
#                     text-align: left;
#                     font-size: 1.2rem;
#                     line-height: 28px;
#                     margin-bottom: 20px;
#                     color: {self.color};
#                     font-weight: bold;
#                     text-transform: capitalize;
#                 }}
#                 .page-number {{
#                     text-align: right;
#                     font-size: 0.9em;
#                     color: {self.color};
#                     text-transform: uppercase;
#                     position: absolute;
#                     bottom: 20px;
#                     right: 20px;
#                 }}
#             </style>
#         </head>
#         <body>
#             {formatted_content}
#         </body>
#         </html>
#         """
#         return self.html_content

#     def save_html(self, output_path="output_report.html"):
#         if not self.html_content:
#             self.format_content_to_html()
#         with open(output_path, "w", encoding="utf-8") as file:
#             file.write(self.html_content)
#         print(f"Relatório formatado e salvo com sucesso como '{output_path}'.")


class BookFormatter:
    def __init__(self, file_path, color="black", font="Courier", 
                 page_background="#FAF0E6", book_background="#D3D3D3", 
                 font_size=16, line_height=1.5, page_width=600, page_height=1000):
        self.file_path = file_path
        self.color = color
        self.font = font
        self.page_background = page_background
        self.book_background = book_background
        self.font_size = font_size
        self.line_height = line_height
        self.page_width = page_width
        self.page_height = page_height
        self.content = ""
        self.html_content = ""
        
    def read_docx(self):
      doc = Document(self.file_path)
      content = ""
      for paragraph in doc.paragraphs:
          paragraph_text = paragraph.text.strip()
          if re.match(r'^(Capítulo|CAPÍTULO)\s+\d+', paragraph_text):
              content += f"<page-breaker> <div class='chapter-title'>{paragraph_text}</div>\n"
          elif re.match(r'^(Prólogo|PRÓLOGO)', paragraph_text):
              content += f"<div class='chapter-title'>{paragraph_text}</div>\n"
          else:
              content += paragraph_text + "\n"  
      self.content = content.strip()  
      return self.content


    def split_content_into_pages(self):
      chars_per_line = int(self.page_width / (self.font_size * 0.85))  
      line_height_px = self.font_size * self.line_height
      lines_per_page = int(self.page_height / line_height_px)
      max_chars_per_page = chars_per_line * lines_per_page

      words = self.content.split()
      pages = []
      current_page = ""
      current_chars = 0

      for word in words:
          print(f'THREATHING WORD BY DEFAULT PROCESS ({word}): {random.getrandbits(128)}')

          if current_chars + len(word) + 1 > max_chars_per_page:
              pages.append(current_page.strip())  
              current_page = ""  
              current_chars = 0
          elif re.match(r'<page-breaker>', word):
              pages.append(current_page.strip())  
              current_page = ""  
              current_chars = 0

          
          current_page += word + " "
          current_chars += len(word) + 1  

      if current_page:  
          pages.append(current_page.strip())

      return pages
  
  
    def format_content_to_html(self):
        pages = self.split_content_into_pages()
        formatted_content = ""

        for page_number, page_content in enumerate(pages, start=1):
            formatted_content += f"""
            <div class="page">
                {page_content.replace("\n", "<br>")}
                <div class="page-number">Página {page_number}</div>
            </div>
            """

        self.html_content = f"""
        <!DOCTYPE html>
        <html lang="en">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>Formatted Book</title>
            <style>
                body {{
                    font-family: {self.font};
                    background-color: {self.book_background};
                    display: flex;
                    justify-content: center;
                    padding: 20px;
                    flex-wrap: wrap;
                }}
                .page {{
                    background-color: {self.page_background};
                    color: {self.color};
                    width: {self.page_width}px;
                    height: {self.page_height}px;
                    padding: 94px 40px 40px;
                    margin: 20px;
                    box-shadow: 2px 2px 5px rgba(0, 0, 0, 0.1);
                    overflow: hidden;
                    position: relative;
                    font-size: {self.font_size}px;
                    line-height: {self.line_height};
                    box-sizing: border-box;
                    position: relative;
                    text-align:justify;
                }}
                .chapter-title {{
                    text-align: center;
                    font-size: 1.3rem;
                    line-height: 30px;
                    margin-bottom: 24px;
                    color: {self.color};
                    font-weight: bold;
                    text-transform: uppercase;
                    position: absolute;
                    top: 40px;
                    left: 50%;
                    transform: translateX(-50%);
                }}
                .page-number {{
                    text-align: center;
                    font-size: 0.9em;
                    color: {self.color};
                    text-transform: uppercase;
                    position: absolute;
                    bottom: 20px;
                    right: 20px;
                }}
            </style>
        </head>
        <body>
            {formatted_content}
        </body>
        </html>
        """
        return self.html_content

    def save_html(self, output_path="output.html"):
        if not self.html_content:
            self.format_content_to_html()
        with open(output_path, "w", encoding="utf-8") as file:
            file.write(self.html_content)
        print(f"HTML formatado e salvo com sucesso como '{output_path}'.")


file_path = "doc.docx"  
formatter = BookFormatter(file_path)
formatter.read_docx()
formatter.format_content_to_html()
formatter.save_html("output.html")
